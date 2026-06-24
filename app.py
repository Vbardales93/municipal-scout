import os
import json
import streamlit as st
from google import genai
from google.genai import types
from pydantic import BaseModel, Field
from typing import List, Literal
from fpdf import FPDF
import pypdf
import docx
from supabase import create_client, Client


# =====================================================================
# 1. ADVANCED PDF REPORT GENERATION ENGINE (Custom Scaffolding Class)
# =====================================================================
class MunicipalScoutReport(FPDF):
    """Custom FPDF engine providing headers, footers, and grid tracking."""

    def __init__(self, town_name, project_type, property_address):
        super().__init__()
        self.town_name = town_name
        self.project_type = project_type
        self.property_address = property_address

    def header(self):
        # Top branding banner
        self.set_fill_color(240, 244, 248)
        self.rect(0, 0, 210, 35, "F")

        # Primary Title Typography
        self.set_font("Helvetica", "B", 16)
        self.set_text_color(20, 40, 80)
        self.set_y(10)
        self.cell(0, 8, f"MUNICIPAL SCOUT: OFFICIAL COMPLIANCE AUDIT", new_x="LMARGIN", new_y="NEXT", align="C")

        # Metadata Subtitle
        self.set_font("Helvetica", "I", 10)
        self.set_text_color(100, 100, 100)
        self.cell(0, 6,
                  f"Jurisdiction: {self.town_name}, CT  |  Project: {self.project_type}  |  Target: {self.property_address}",
                  new_x="LMARGIN", new_y="NEXT", align="C")

        # Decorative Separator Line
        self.set_draw_color(20, 40, 80)
        self.set_line_width(0.5)
        self.line(10, 30, 200, 30)
        self.ln(12)

    def footer(self):
        # Position footer at 15 mm from bottom
        self.set_y(-15)
        self.set_draw_color(220, 224, 230)
        self.line(10, self.get_y() - 2, 200, self.get_y() - 2)

        # Page Numbers & Legal Disclaimer
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(150, 150, 150)
        self.cell(0, 10,
                  "Disclaimer: This AI-generated audit is a pre-flight assessment. Final validation rests with local building officials.",
                  new_x="RIGHT", new_y="KEEP", align="L")
        self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", new_x="LMARGIN", new_y="NEXT", align="R")


# =====================================================================
# 2. DATA SCHEMAS FOR STRUCTURED AI OUTPUT (Dynamic Checklist Style)
# =====================================================================
class RuleEvaluation(BaseModel):
    rule_section: str = Field(
        description="The section or clause number discovered in the PDF (e.g., R326.6.1.1 or Local Section 12-A)")
    rule_title: str = Field(description="A short, descriptive title of the rule requirement")
    status: Literal["PASS", "VIOLATION"] = Field(
        description="Mark PASS if the proposal fully satisfies this rule. Mark VIOLATION if the proposal violates it or completely omits mandatory details.")
    exact_rule_text: str = Field(description="The verbatim or highly precise text of the rule from the PDF.")
    precise_correction_needed: str = Field(
        description="Detailed instructions for the builder on how to fix this if it is a VIOLATION. Leave blank if PASS.")


class ComprehensiveAudit(BaseModel):
    summary: str = Field(
        description="A professional, introductory summary assessment written as a Municipal Zoning Officer.")
    all_evaluations: List[RuleEvaluation] = Field(
        description="A complete checklist of EVERY single distinct rule requirement found in the rules PDF.")


# =====================================================================
# 3. ADVANCED TEXT EXTRACTION HELPERS (PDF Layers & Word Layout Tables)
# =====================================================================
def extract_text_from_docx(uploaded_file):
    """Extracts text from both standard paragraphs and layout tables inside Word documents."""
    doc = docx.Document(uploaded_file)
    full_text = []

    # 1. Ingest baseline body paragraph strings
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            full_text.append(paragraph.text)

    # 2. Ingest table matrices cell by cell, preserving structural rows
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                # Deduplicate cell values (handles merged document cells smoothly)
                clean_row = []
                for item in row_text:
                    if not clean_row or item != clean_row[-1]:
                        clean_row.append(item)
                full_text.append(" | ".join(clean_row))

    return "\n".join(full_text)


# =====================================================================
# 4. INITIALIZE SECURE CLOUD DATABASE (Supabase Connectivity)
# =====================================================================
try:
    supabase_url = st.secrets["SUPABASE_URL"]
    supabase_key = st.secrets["SUPABASE_KEY"]
    supabase: Client = create_client(supabase_url, supabase_key)
except Exception as e:
    st.error(f"⚠️ Cloud Database Connection Failure: Check your production secrets file. Details: {e}")

# =====================================================================
# 5. STREAMLIT UI VIEW SCRIPTS (Centered Layout Setup)
# =====================================================================
st.set_page_config(page_title="Municipal Scout", page_icon="🔍", layout="centered")

# Global UI Style Injections
st.markdown("""
<style>
    .report-metric { font-size: 24px; font-weight: bold; color: #142850; }
    .stButton>button { width: 100%; border-radius: 6px; }
</style>
""", unsafe_allow_html=True)

# Main Navigation Tabs separating the interface operations cleanly
workspace_tab, history_tab = st.tabs(["🔍 Audit Production Workspace", "🕒 Historical Cloud Logs"])

# Side-panel parameters shared across global views
st.sidebar.header("Audit Parametrics")
selected_town = st.sidebar.selectbox("Target Jurisdiction", ["Windsor Locks", "Bristol"])
selected_project = st.sidebar.selectbox("Regulatory Scope", ["Pool", "Deck", "Shed"])
target_address = st.sidebar.text_input("Property Site Address", "74 Center Street")
api_key_input = st.sidebar.text_input("Gemini Production API Key", type="password")

# =====================================================================
# VIEW A: THE PRIMARY WORKSPACE TAB
# =====================================================================
with workspace_tab:
    st.markdown("<h1 style='text-align: center;'>🔍 Municipal Scout</h1>", unsafe_allow_html=True)
    st.markdown("<h4 style='text-align: center; color: gray;'>Pre-Flight Hyper-Local Zoning & Permit Auditor</h4>",
                unsafe_allow_html=True)
    st.markdown("---")

    st.write("### 📄 Step 1: Upload Project Outline / Proposal")
    uploaded_file = st.file_uploader(
        f"Ingest contractor {selected_project.lower()} blueprint text layers (.txt, .pdf, .docx)",
        type=["txt", "pdf", "docx"]
    )

    if uploaded_file is not None:
        if uploaded_file.name.endswith(".pdf"):
            with st.spinner("Deconstructing PDF vector text elements..."):
                pdf_reader = pypdf.PdfReader(uploaded_file)
                application_content = ""
                for page in pdf_reader.pages:
                    text_layer = page.extract_text()
                    if text_layer:
                        application_content += text_layer + "\n"
            st.success(f"Extracted {len(pdf_reader.pages)} PDF pages successfully.")

        elif uploaded_file.name.endswith(".docx"):
            with st.spinner("Extracting content blocks and data layouts from Word file..."):
                application_content = extract_text_from_docx(uploaded_file)
            st.success("Extracted text layers and structural data tables successfully.")

        else:
            application_content = uploaded_file.read().decode("utf-8")
            st.success("Flat configuration text loaded successfully.")

        with st.expander("👀 View Extracted Document Content Payload"):
            st.text(application_content)

        st.write("### ⚡ Step 2: Run Compliance Audit")
        if st.button("🚀 Launch Compliance Audit"):
            if not api_key_input:
                st.error("⚠️ System halted: Missing Gemini API verification key in configuration sidebar.")
                st.stop()

            # Map dynamic file paths based on chosen municipal parameters
            town_folder = selected_town.lower().replace(" ", "_")
            project_file = selected_project.lower()
            rule_path = f"rules/{town_folder}/{project_file}.pdf"

            if not os.path.exists(rule_path):
                st.error(f"⚠️ Missing Local Rulebook: Reference document '{rule_path}' could not be located on disk.")
                st.stop()

            with open(rule_path, "rb") as file:
                rules_pdf_bytes = file.read()

            with st.spinner(f"Ingesting {selected_town} codebook and running dual-task compliance evaluation..."):
                try:
                    # Instantiate Google GenAI Client
                    client = genai.Client(api_key=api_key_input)

                    system_instruction = (
                        f"You are a highly meticulous Municipal Zoning Enforcement Officer and Building Inspector in Connecticut. "
                        f"Your job is to perform a systematic, two-step audit on the attached residential {selected_project.lower()} proposal using the provided rules PDF.\n\n"
                        f"DYNAMIC AUDIT PROTOCOL:\n"
                        f"1. GENERATE THE CHECKLIST: Read the rules PDF and dynamically extract EVERY single individual requirement, parameter, threshold, and local rule. This is your temporary master checklist for this town.\n"
                        f"2. EVALUATE EVERYTHING: Evaluate the contractor's proposal against EVERY single rule you just extracted. Do not skip any.\n"
                        f"3. REPORT ALL RULES: Your output must include an entry in 'all_evaluations' for every single rule discovered. If the proposal satisfies the rule, mark it PASS. If the proposal fails the rule OR completely omits a mandatory detail (such as missing a required specification), mark it VIOLATION.\n"
                        f"4. CRITICAL: Never bundle separate rules together. Keep your discovered checklist completely atomic to guarantee identical item counts across separate runs of this same PDF."
                    )

                    user_prompt = f"""You are auditing this residential {selected_project.lower()} proposal for the municipality of {selected_town}, Connecticut. 

CRITICAL: Ignore any town names, cities, or zip codes written inside the contractor's proposal text. Assume this proposal has been officially submitted to the {selected_town} Building and Zoning Department. 

Audit the provided contractor proposal text strictly against the attached official town {selected_project.lower()} rules PDF document.

=== CONTRACTOR PROPOSAL ===
{application_content}"""

                    # Query Structured Model Pipeline
                    response = client.models.generate_content(
                        model='gemini-2.5-flash',
                        contents=[
                            types.Part.from_bytes(data=rules_pdf_bytes, mime_type='application/pdf'),
                            user_prompt
                        ],
                        config=types.GenerateContentConfig(
                            system_instruction=system_instruction,
                            temperature=0.0,
                            response_mime_type="application/json",
                            response_schema=ComprehensiveAudit,
                        )
                    )

                    if response.parsed:
                        audit_data = response.parsed.model_dump()
                    else:
                        st.error("⚠️ Critical API Error: Empty structural data model returned.")
                        st.stop()

                    # Process data sets into logical validation buckets
                    all_items = audit_data.get("all_evaluations", [])
                    violations = [item for item in all_items if item["status"] == "VIOLATION"]
                    passes = [item for item in all_items if item["status"] == "PASS"]
                    total_issues = len(violations)

                    # Render Plain-Text compilation for report pipelines
                    report_text = f"{audit_data.get('summary', '')}\n\n"
                    report_text += f"=== IDENTIFIED CODE VIOLATIONS & OMISSIONS ({total_issues} Total) ===\n\n"
                    for idx, item in enumerate(violations, 1):
                        report_text += f"{idx}. {item['rule_title']} ({item['rule_section']})\n"
                        report_text += f"- Rule text: {item['exact_rule_text']}\n"
                        report_text += f"- Required Correction: {item['precise_correction_needed']}\n\n"

                    # --- DB TRANSACTION ---
                    try:
                        supabase.table("audits").insert({
                            "property_address": target_address,
                            "municipality": selected_town,
                            "project_category": selected_project,
                            "issues_found": total_issues,
                            "detailed_report": report_text
                        }).execute()
                    except Exception as db_err:
                        st.warning(f"⚠️ Logging bypass: Metrics could not be sent to Supabase storage. ({db_err})")

                    # --- WORKSPACE DASHBOARD DISPLAY ---
                    st.markdown("---")
                    st.write("### 📊 Live Audit Metrics")

                    metric_col1, metric_col2, metric_col3 = st.columns(3)
                    with metric_col1:
                        st.metric(label="Total Criteria Evaluated", value=len(all_items))
                    with metric_col2:
                        st.metric(label="Compliant Criteria Pass", value=len(passes))
                    with metric_col3:
                        st.metric(label="Non-Compliance Violations", value=total_issues)

                    st.subheader("📝 Executive Legal Summary")
                    st.info(audit_data.get("summary", "No summary text generated."))

                    st.subheader("🚨 Code Violations & Missing Omissions Detailed Matrix")
                    if total_issues == 0:
                        st.success("🎉 Compliance verification complete. No anomalies detected against local codebooks.")
                    else:
                        for idx, v in enumerate(violations, 1):
                            st.warning(f"**{idx}. {v['rule_title']} ({v['rule_section']})**")
                            st.markdown(f"- **Official Rule Text:** *{v['exact_rule_text']}*")
                            st.markdown(f"- **Actionable Builder Correction:** {v['precise_correction_needed']}")
                            st.markdown("---")

                    with st.expander("✅ View Passed Checklist Items"):
                        for p in passes:
                            st.markdown(f"**🟢 {p['rule_title']} ({p['rule_section']})** — *Status: Verified Compliant*")
                            st.caption(f"Code Context: {p['exact_rule_text']}")

                    # --- CUSTOM COMPATIBLE PDF COMPILATION ---
                    pdf = MunicipalScoutReport(selected_town, selected_project, target_address)
                    pdf.set_auto_page_break(auto=True, margin=15)
                    pdf.add_page()
                    pdf.set_font("Helvetica", "", 11)

                    # Standardize custom windows characters to standard Latin-1 strings safely
                    sanitized_pdf_string = (
                        report_text
                        .replace('“', '"').replace('”', '"')
                        .replace('’', "'").replace('‘', "'")
                        .replace('—', '-').replace('–', '-')
                    )
                    pdf.multi_cell(0, 6, sanitized_pdf_string)

                    raw_binary_output = pdf.output(dest='S')
                    compiled_bytes = bytes(raw_binary_output) if not isinstance(raw_binary_output,
                                                                                str) else raw_binary_output.encode(
                        'latin-1')

                    st.markdown("---")
                    st.write("### 💾 Step 3: Export Verified Compliance Records")
                    export_name = f"{selected_town.lower().replace(' ', '_')}_{project_file}_audit_report.pdf"

                    st.download_button(
                        label="📥 Download Printable PDF Report",
                        data=compiled_bytes,
                        file_name=export_name,
                        mime="application/pdf"
                    )

                except Exception as api_err:
                    if "429" in str(api_err) or "RESOURCE_EXHAUSTED" in str(api_err):
                        st.error(
                            "⏳ **Gemini API Cooldown Active:** Free tier transaction thresholds breached. Wait approximately 45 seconds before refiring.")
                    else:
                        st.error(f"Processing Exception Intercepted: {api_err}")

# =====================================================================
# VIEW B: INTERACTIVE DEEP HISTORY EXPLORER TAB (Supabase Framework)
# =====================================================================
with history_tab:
    st.markdown("<h2 style='text-align: center;'>🕒 Cloud Record History Explorer</h2>", unsafe_allow_html=True)
    st.write(
        "Query, read, and inspect full plaintext audit briefs generated previously across secure database environments.")
    st.markdown("---")

    try:
        # Pull transactional history files
        db_fetch = supabase.table("audits").select(
            "id, created_at, property_address, municipality, project_category, issues_found, detailed_report").order(
            "created_at", desc=True).execute()

        if db_fetch.data:
            st.write("### 🔍 Deep Inspection: Read Full Past Audit Text Block")

            # Map descriptive titles to full database records dynamically
            record_map = {
                f"ID {row['id']}: {row['property_address']} ({row['municipality']} - {row['project_category']})": row
                for row in db_fetch.data
            }

            selected_history_key = st.selectbox("Select a past log entry to view full multi-line details:",
                                                list(record_map.keys()), key="history_tab_selector")

            if selected_history_key:
                chosen_record = record_map[selected_history_key]
                st.markdown(f"#### 📄 Full Technical Briefing Log: {chosen_record['property_address']}")
                st.caption(f"Database Record Created: {chosen_record['created_at']}")

                # Render the text box with collapsed labels for a streamlined user experience
                st.text_area("", value=chosen_record["detailed_report"], height=350, key="history_tab_text_area",
                             label_visibility="collapsed")

        else:
            st.info("Cloud storage registers are empty. Run your initial evaluation workflow to seed history items.")

    except Exception as history_error:
        st.caption(f"SQL Storage Feed Engine Offline: Could not map database table vectors. ({history_error})")